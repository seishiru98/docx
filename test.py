import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT

def read_excel_data(filename, sheet_name):
    try:
        return pd.read_excel(filename, sheet_name=sheet_name)
    except FileNotFoundError:
        print(f"ОШИБКА: Файл {filename} не найден.")
    except ValueError:
        print(f"ОШИБКА: Лист {sheet_name} не найден в файле {filename}.")
    except Exception as e:
        print(f"ОШИБКА: Произошла ошибка при чтении файла {filename}: {e}")

def add_header(doc, header_text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(header_text)
    run.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_table(doc, df, start_row, end_row):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Adding header row
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        cell_paragraph = hdr_cells[i].paragraphs[0]
        cell_paragraph.text = column_name
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in cell_paragraph.runs:
            set_font(run, 'Times New Roman', 12)
        set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                             line_spacing=18, space_after=0, space_before=0)

    # Adding data rows
    for index in range(start_row, end_row):
        row = df.iloc[index]
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_paragraph = row_cells[i].paragraphs[0]
            cell_paragraph.text = str(value)
            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in cell_paragraph.runs:
                set_font(run, 'Times New Roman', 12)
            set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                                 line_spacing=18, space_after=0, space_before=0)

def set_font(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_paragraph_format(paragraph, left_indent, right_indent, first_line_indent, line_spacing, space_after, space_before):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Pt(left_indent)
    paragraph_format.right_indent = Pt(right_indent)
    paragraph_format.first_line_indent = Pt(first_line_indent)
    paragraph_format.line_spacing = Pt(line_spacing)
    paragraph_format.space_after = Pt(space_after)
    paragraph_format.space_before = Pt(space_before)

def insert_page_break(doc):
    doc.add_page_break()

# Main code
df7_1 = read_excel_data('database.xlsx', '7.1')
df7_1 = df7_1.fillna('')

doc = Document()
table_name = '7.1'
header_text = f'Таблица {table_name} – Эксплуатационные расходы энергоресурсов при демеркаптанизации СУГ'

rows_per_page = 13  # Adjust this number based on your page size and table formatting

total_rows = len(df7_1)
start_row = 0

while start_row < total_rows:
    end_row = min(start_row + rows_per_page, total_rows)
    add_header(doc, header_text)
    add_table(doc, df7_1, start_row, end_row)
    start_row = end_row
    if start_row < total_rows:
        insert_page_break(doc)


