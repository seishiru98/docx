from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
import pandas as pd

# Создаем новый документ
doc = Document()

section = doc.sections[0]

# Установка полей страницы
section.left_margin = Cm(2)  # Левое поле
section.right_margin = Cm(1)  # Правое поле
section.top_margin = Cm(2)  # Верхнее поле
section.bottom_margin = Cm(2)  # Нижнее поле

# Функция для изменения шрифта и размера шрифта
def set_font(run, font_name, font_size, italic=False, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.italic = italic
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # Это нужно для корректного отображения шрифта на всех платформах
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)


# Функция для установки отступов и междустрочного интервала
def set_paragraph_format(paragraph, left_indent=0, right_indent=0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Cm(left_indent)
    paragraph_format.right_indent = Cm(right_indent)
    paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph_format.line_spacing = Pt(line_spacing)
    paragraph_format.space_after = Cm(space_after)
    paragraph_format.space_before = Cm(space_before)

def add_header(doc, header_text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(header_text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def add_table(doc, df, start_row, end_row):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        cell_paragraph = hdr_cells[i].paragraphs[0]
        cell_paragraph.text = column_name
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in cell_paragraph.runs:
            set_font(run, 'Times New Roman', 12)
        set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                             line_spacing=18, space_after=0, space_before=0)

    for index in range(start_row, end_row):
        row = df.iloc[index]
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_paragraph = row_cells[i].paragraphs[0]
            cell_paragraph.text = str(value)
            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in cell_paragraph.runs:
                set_font(run, 'Times New Roman', 12)
            set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                                 line_spacing=18, space_after=0, space_before=0)

        for col_idx in range(len(df.columns)):
            merge_start = 2
            for row_idx in range(2, len(table.rows) + 1):
                cell = table.cell(row_idx - 1, col_idx)
                if cell.text == "":
                    continue
                if merge_start < row_idx - 1:
                    table.cell(merge_start - 1, col_idx).merge(table.cell(row_idx - 2, col_idx))
                merge_start = row_idx
            if merge_start < len(table.rows):
                table.cell(merge_start - 1, col_idx).merge(table.cell(len(table.rows) - 1, col_idx))

def insert_page_break(doc):
    doc.add_page_break()

class Counter:
    def __init__(self, start_value, step):
        self.value = start_value
        self.step = step

    def increment(self):
        current_value = self.value
        self.value += self.step
        return current_value

class HeadingCounter(Counter):
    def __init__(self, start_value, paragraph_counter, table_counter, fig_counter):
        super().__init__(start_value, 1)
        self.paragraph_counter = paragraph_counter
        self.table_counter = table_counter
        self.fig_counter = fig_counter

    def increment(self):
        current_value = super().increment()
        self.paragraph_counter.reset(current_value + 0.1)
        self.table_counter.reset(current_value + 0.1)
        self.fig_counter.reset(current_value + 0.1)
        return current_value

class ParagraphCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value

class TableCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value

class FigCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value

# Инициализация счетчиков
n_heading = 1
n_paragraph_start = n_heading + 0.1
n_table_start = n_heading + 0.1
n_fig_start = n_heading + 0.1

par_counter = ParagraphCounter(n_paragraph_start, 0.1)
table_counter = TableCounter(n_table_start, 0.1)
fig_counter = FigCounter(n_fig_start, 0.1)
head_counter = HeadingCounter(n_heading, par_counter, table_counter, fig_counter)

def read_excel_data(filename, sheet_name):
    try:
        return pd.read_excel(filename, sheet_name=sheet_name)
    except FileNotFoundError:
        print(f"ОШИБКА: Файл {filename} не найден.")
    except ValueError:
        print(f"ОШИБКА: Лист {sheet_name} не найден в файле {filename}.")
    except Exception as e:
        print(f"ОШИБКА: Произошла ошибка при чтении файла {filename}: {e}")

#-----------------------------------------------------------------------------------------------------------------------

text = ['ООО «НТЦ «Ахмадуллины»',
        '']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 18)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = ['УТВЕРЖДАЮ',
        'Генеральный директор',
        'Р.М. Ахмадуллин',
       f'«{None}» {None} {None} года',
        '']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = ['Исходные данные для проектирования',
        'блока демеркаптанизации СУГ установки замедленного коксования Комплекса глубокой переработки нефти',
        '(базовый проект)',
        '']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14, bold=True)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = ['198-22-001.001.009-ИД']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14, bold=True)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = ['В настоящем документе содержится конфиденциальная информация относительно технологии «Демерус», включая эксплуатационные условия и технологические возможности, которые не могут быть раскрыты неуполномоченным лицам. Представленные материалы являются собственностью Лицензиара. Получая настоящую информацию, вы соглашаетесь не использовать ее ни для каких других целей, кроме тех, которые согласованы с Лицензиаром в письменной форме, не воспроизводить этот документ полностью или частично и не раскрывать его содержимое третьим лицам без письменного разрешения Лицензиара.',
        '',
        '',
        '',
        '']

# Добавление абзацев и установка их форматирования
for line in text:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = paragraph.add_run(line)
    set_font(run, 'Times New Roman', 12, italic=True)  # Установить курсив
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=18,
                         space_after=0, space_before=0)

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'

header_text = ['№ п/п', 'Ревизия', 'Дата выдачи']

hdr_cells = table.rows[0].cells
for i, text in enumerate(header_text):
    cell_paragraph = hdr_cells[i].paragraphs[0]
    cell_paragraph.text = text
    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in cell_paragraph.runs:
        set_font(run, 'Times New Roman', 14, bold=True)
    set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                         line_spacing=22, space_after=0, space_before=0)

text = ['',
        '',
        '',
        'Казань 2024']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

ch_1 = head_counter.increment()

heading = doc.add_heading(f'{ch_1:.0f} ВВЕДЕНИЕ', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = ['',
        '',
        'Настоящий документ разработан на основании договора № 13-2/22 от 10.10.2022 г. в соответствии с Техническим Заданием на разработку Базового проекта очистки сжиженных углеводородных газов от меркаптанов АО «Инженерно-промышленная нефтехимическая компания» для ПАО «Славнефть-ЯНОС».'
        'Блок щелочной очистки СУГ предназначен для удаления меркаптанов и остаточного сероводорода из СУГ и рассчитан на переработку по номинальной производительности 16210,0 кг/ч СУГ.',
        'В состав блока щелочной очистки «ДЕМЕРУС» входят:',
        '– узел очистки СУГ от меркаптановой серы и остаточного сероводорода водным раствором гидроксида натрия;',
        '– окислительно-каталитическая регенерация щелочного раствора;',
        '– реагентное хозяйство (емкость хранения и приготовления щелочного раствора V-305).',
        'Разработчик Базового проекта блока щелочной очистки СУГ «ДЕМЕРУС» с регенерацией отработанного раствора щелочи – ООО «НТЦ «Ахмадуллины».',
        'Режим работы блока щелочной очистки СУГ «ДЕМЕРУС» с регенерацией отработанного раствора щелочи – круглосуточный, круглогодичный 8760 часов в год. Расчетный период непрерывной эксплуатации установки между остановками на капитальный ремонт – 48 месяцев. Срок службы оборудования не менее 20 лет. Срок службы катализатора не менее 8 лет. При расчете и подборе оборудования, согласно Технического задания на проектирование, был принят диапазон устойчивой производительности 60÷110% от расчетного расхода.',
        '']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

#-----------------------------------------------------------------------------------------------------------------------
# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

ch_2 = head_counter.increment()

heading = doc.add_heading(f'{ch_2:.0f} ОБЩИЕ СВЕДЕНИЯ О ТЕХНОЛОГИИ «ДЕМЕРУС»', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

ch_2_par_1 = par_counter.increment()

text = ['',
        '',
        'Из известных методов очистки СУГ от меркаптанов широко используемым является метод их экстракции щелочным раствором из газов с последующей регенерацией щелочного раствора окислением меркаптидов кислородом воздуха в органические дисульфиды в присутствии гомогенных [1, 2] или гетерогенных [3-5] катализаторов.',
        f'Содержание общей серы в СУГ после демеркаптанизации на гомогенных и гетерогенных катализаторах представлено в таблице {ch_2_par_1}',
        ]

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

df2_1 = read_excel_data('database.xlsx', '2.1')
df2_1 = df2_1.fillna('')

header_text_first = f'Таблица {ch_2_par_1:.1f} – Содержание общей серы в СУГ после демеркаптанизации на гомогенных и гетерогенных катализаторах'
header_text_next = f'Продолжение таблицы {ch_2_par_1:.1f} – Содержание общей серы в СУГ после демеркаптанизации на гомогенных и гетерогенных катализаторах'

rows_per_page_first = 9  # Количество строк для первой таблицы
rows_per_page_next = 12  # Количество строк для следующих таблиц

total_rows = len(df2_1)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df2_1, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df2_1, start_row, end_row)
    start_row = end_row

text = ['',
        'При использовании гомогенного (растворенного в щелочи) фталоцианинового катализатора процесс окисления меркаптидов продолжается и вне регенератора: в трубопроводах и в экстракторе по причине совместного присутствия в циркулирующем щелочном растворе катализатора и растворенного кислорода. Образующиеся вне регенератора дисульфиды переходят из щелочи в очищаемый углеводород в процессе экстракции СУГ, повышая в нем содержание общей серы и жидкого остатка [1, 5], что является большим недостатком гомогенно-каталитических процессов демеркаптанизации СУГ, особенно при использовании СУГ в качестве моторного, коммунально-бытового топлива, изобутиленовой фракции - сырья МТБЭ, пропиленовой фракции - для полимеризации.',
        'Технология демеркаптанизации углеводородного сырья «Демерус» разработана НТЦ «AhmadullinS» в 90-х годах XX века для решения многоплановых задач удаления меркаптанов — демеркаптанизации природного газа, сжиженных углеводородных газов, бензиновых, керосиновых и дизельных фракций, газовых конденсатов и легких нефтей.',
        'Отличительной особенностью технологии «Демерус» по сравнению с отечественными и зарубежными аналогами является использование при регенерации меркаптидсодержащей щелочи гетерогенного фталоцианинового катализатора КСМ-Х, изготовленного нанесением каталитически активных компонентов на полимерную основу.',
        'В отличие от водорастворимых фталоцианиновых катализаторов, подверженных термическому и гидролитическому разложению в водно-щелочных растворах, состав и технология производства гетерогенного катализатора КСМ-Х, обеспечивает прочное удерживание его каталитически активных компонентов на полимерном носителе. Это исключает необходимость периодической или непрерывной подпитки катализатора КСМ-Х дорогостоящими соединениями металлов переменной валентности, т.е. их нежелательное расходование и загрязнение сточных вод предприятия солями тяжелых металлов. Повышенная стойкость этого катализатора к каталитическим ядам и термическому воздействию обеспечивает его стабильную активность на протяжении всего срока промышленной эксплуатации.',
        'Многолетний опыт промышленной эксплуатации катализатора КСМ-Х показал следующие преимущества по сравнению с гомогенными катализаторами [3-6]:',
        '1. Срок эксплуатации катализатора в системе очистки возрастает ≈ в 30 раз (с 3÷4-х месяцев ≈ до 10 лет).',
        '2. Срок службы щелочного раствора увеличивается примерно втрое и составляет 1 год (без замены), что позволяет значительно сократить расход щелочи и объем щелочных стоков с блоков демеркаптанизации углеводородных газов.',
        '3. Использование гетерогенного катализатора КСМ-Х позволяет исключить попадание солей тяжелых металлов - фталоцианина кобальта и его производных в сточные воды и далее на БОС и в водоемы.',
        'Особенностью предлагаемого способа регенерации щелочи для очистки СУГ по сравнению с известными отечественными и зарубежными аналогами является использование при регенерации меркаптидной щелочи усовершенствованного гетерогенного фталоцианинового катализатора серии КСМ-Х [7], преимуществом которого является его устойчивость к воздействию примесей аминов в щелочном растворе [8].',
        ]

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)


#-----------------------------------------------------------------------------------------------------------------------
# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width


ch_3 = head_counter.increment()

heading = doc.add_heading(f'{ch_3:.0f} ТЕХНИЧЕСКИЙ УРОВЕНЬ, ПАТЕНТОСПОСОБНОСТЬ И ПАТЕНТНАЯ ЧИСТОТА ПРОЦЕССА', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

ch_3_par_1 = par_counter.increment()
ch_3_par_2 = par_counter.increment()

text = ['',
        '',
        'Результаты обследования и анализ работы существующих отечественных и зарубежных процессов демеркаптанизации легкого углеводородного сырья показали, что для регенерации щелочи эффективным является метод окисления сернистых соединений кислородом воздуха в присутствии гомогенного [1] или гетерогенного [2-5] фталоцианиновых катализаторов. ',
        'Собственники ООО «НТЦ «Ахмадуллины» являются Лицензиарами технологии «ДЕМЕРУС» (патент № 2603635), применяемой для демеркаптанизации СУГ.',
        'В технологии «ДЕМЕРУС» применяется катализатор КСМ-Х, изготавливаемый по ТУ 2175-001-40655797-2014, состав и технология приготовления которого защищены патентами РФ на изобретение № 2529500 и № 2677226. Гарантийный срок эксплуатации катализатора КСМ-Х составляет 8 лет.',
        'Катализатор КСМ-Х изготовлен нанесением каталитически активных компонентов на полимерный носитель и выполнен в виде блочных насадочных элементов с развитой геометрической поверхностью. Катализатор стационарно закреплен в регенераторе. Состав и технология его изготовления исключают вымывание каталитически активных компонентов из полимерного носителя катализатора и попадание их в щелочной раствор, что обеспечивает сферу действия катализатора КСМ-Х только в объеме регенератора. Данные работы [9] показывают, что степень окисления меркаптидов кислородом воздуха в щелочной среде в отсутствие катализаторов чрезвычайно низкая.',
        'Кроме того, процесс окислительной регенерации меркаптидсодержащей щелочи на этом катализаторе можно проводить при повышенных температурах порядка 60-800С, при которых концентрация растворенного кислорода в циркулирующем щелочном растворе примерно вдвое ниже, чем в гомогенно-каталитическом процессе [1], проводимом при 40-500С (из-за низкой термогидролитической устойчивости гомогенного катализатора в щелочи).',
        'Гетерогенный катализатор КСМ-Х хорошо зарекомендовал себя в процессе демеркаптанизации СУГ на 9-ти нефтеперерабатывающих заводах России и ближнего Зарубежья (ОАО АНК "БАШНЕФТЬ"; ОАО "Газпромнефть-МНПЗ" - 3 установки; ООО «ЛУКОЙЛ-Нижегороднефтеоргсинтез» - 2 установки; ОАО "Славнефть-ЯНОС"; ОАО «ТАИФ-НК»; НК «РОСНЕФТЬ» Лисичанский НПЗ (Украина); ORLEN Lietuva (Мажейкяйский НПЗ), Литва; ОАО «Мозырский НПЗ», Беларусь; ООО «ЛУКОЙЛ-Ухтанефтепереработка») [8].',
        'В настоящее время ведется детальное проектирование и строительство еще 8-ми установок сероочистки СУГ, включающий узел регенерации насыщенной меркаптидами щелочи с использованием гетерогенного катализатора КСМ-Х:',
        '1. ООО «Славянск ЭКО» - блок демеркаптанизации сжиженных углеводородных газов с ГФУ (25,9 м3/ч) по технологии «Демерус» с использованием гетерогенного катализатора КСМ-Х.',
        '2.	ОАО «Новошахтинский завод нефтепродуктов» - блок демеркаптанизации сжиженных углеводородных газов с ГФУ (14,8 м3/ч) по технологии «Демерус» с использованием гетерогенного катализатора КСМ-Х.',
        '3.	ООО «Афипский НПЗ» – блок демеркаптанизации сжиженных углеводородных газов с УЗК (13,7 м3/ч) по технологии «Демерус» с использованием гетерогенного катализатора КСМ-Х',
        '4. ОАО «Новошахтинский завод нефтепродуктов» - блок демеркаптанизации сжиженных углеводородных газов с ГФУ (9,7 м3/ч) по технологии «Демерус» с использованием гетерогенного катализатора КСМ-Х.',
        '5. ПАО «Орскнефтеоргсинтез» - блок демеркаптанизации сжиженных углеводо-родных газов с установки замедленного коксования (10,2 м3/ч) и поток предельных СУГ с ГФУ (34,2 м3/ч) по технологии «Демерус» с использованием гетерогенного катализатора КСМ-Х.',
        '6. ООО «Афипский НПЗ» – блок аминовой очистки от сероводорода и щелочной демеркаптанизации сжиженных углеводородных газов (50,0 м3/ч) с установок газофракционирования предельной секции и установки Коксования непредельной секции по технологии «Демерус» с использованием гетерогенного катализатора КСМ-Х.',
        '7. ОАО «РОСНЕФТЬ-Сызранский НПЗ» – блок демеркаптанизации предельного (18,4 м3/ч) и непредельного (33,2 м3/ч) сырья ГФУ по технологии «Демерус» с использованием гетерогенного катализатора КСМ-Х.',
        '8. ОАО АНК «БАШНЕФТЬ» – блок сероочистки (от карбонилсульфида и меркаптанов) сжиженных углеводородных газов (53,0 м3/ч) установки замедленного коксования 2000 по технологии «Демерус» с использованием гетерогенного катализатора КСМ-Х.',
        f'Для сероочистки сжиженных углеводородных газов АО «Газпромнефть-МНПЗ» разработаны следующие технических решения, защищенные патентами РФ (табл. {ch_3_par_1} и {ch_3_par_2}).',
        '',
        ]

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)


df3_1 = read_excel_data('database.xlsx', '3.1')
df3_1 = df3_1.fillna('')

header_text_first = f'Таблица {ch_3_par_1:.1f} – Охранные документы, под действие которых попадает объект техники'
header_text_next = f'Продолжение таблицы {ch_3_par_1:.1f} – Охранные документы, под действие которых попадает объект техники'

rows_per_page_first = 15  # Количество строк для первой таблицы
rows_per_page_next = 15  # Количество строк для следующих таблиц

total_rows = len(df3_1)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df3_1, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df3_1, start_row, end_row)
    start_row = end_row

df3_2 = read_excel_data('database.xlsx', '3.2')
df3_2 = df3_2.fillna('')

header_text_first = f'Таблица {ch_3_par_2:.1f} – Правовая охрана объекта техники (лицензионные договоры)'
header_text_next = f'Продолжение таблицы {ch_3_par_2:.1f} – Правовая охрана объекта техники (лицензионные договоры)'

rows_per_page_first = 15  # Количество строк для первой таблицы
rows_per_page_next = 15  # Количество строк для следующих таблиц

total_rows = len(df3_2)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df3_2, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df3_2, start_row, end_row)
    start_row = end_row

#-----------------------------------------------------------------------------------------------------------------------
# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width


ch_4 = head_counter.increment()

heading = doc.add_heading(f'{ch_4:.0f} ХАРАКТЕРИСТИКА ИСХОДНОГО СЫРЬЯ, ПРОДУКТОВ, ОСНОВНЫХ И ВСПОМОГАТЕЛЬНЫХ МАТЕРИАЛОВ', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

df2_1 = read_excel_data('database.xlsx', 'Лист5')
df2_1 = df2_1.fillna('')

header_text_first = f'Таблица {ch_2_par_1:.1f} – Содержание общей серы в СУГ после демеркаптанизации на гомогенных и гетерогенных катализаторах'
header_text_next = f'Продолжение таблицы {ch_2_par_1:.1f} – Содержание общей серы в СУГ после демеркаптанизации на гомогенных и гетерогенных катализаторах'

rows_per_page_first = 9  # Количество строк для первой таблицы
rows_per_page_next = 12  # Количество строк для следующих таблиц

total_rows = len(df2_1)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df2_1, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df2_1, start_row, end_row)
    start_row = end_row

#-----------------------------------------------------------------------------------------------------------------------
# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width


ch_5 = head_counter.increment()

heading = doc.add_heading(f'{ch_5:.0f} ТЕХНИЧЕСКАЯ ХАРАКТЕРИСТИКА ОТХОДОВ И ОТРАБОТАННОГО ВОЗДУХА', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

#-----------------------------------------------------------------------------------------------------------------------
new_section = doc.add_section(WD_SECTION.NEW_PAGE)
new_section.orientation = WD_ORIENT.PORTRAIT

ch_6 = head_counter.increment()

heading = doc.add_heading(f'{ch_6:.0f} ТЕХНОЛОГИЯ ПРОЦЕССА', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

ch_6_par_1 = par_counter.increment()
fig6_1 = fig_counter.increment()

content1 = [
    "",
    "",
    f"Принципиальная технологическая схема демеркаптанизации СУГ приведена на рис. {fig6_1}",
    f"Исходное сырье с содержанием сероводорода до {None} % мас.; диоксида углерода до {None} % мас., карбонилсульфида по сере до {None} % мас., метилмеркаптана по сере до {None} % мас., этилмеркаптана по сере до {None} % мас. и пропилмеркаптана по сере до {None} % мас. подается в куб насадочного экстрактора {None}. Расход СУГ в экстрактор составляет до {None} кг/ч, температура 40ºС с давлением 21,0 кгс/см2 (изб.). В среднюю часть экстрактора С-301 поверх насадок подается регенерированный водный раствор щелочи с температурой 30÷45°С из куба отстойника дисульфидов V-303. В процессе взаимодействия СУГ со стекающим вниз щелочным раствором происходит хемосорбция содержащихся в нем меркаптанов по реакции 1, щелочной гидролиз карбонилсульфида на 40÷60% по реакции 2, а также хемосорбция сероводорода и диоксида углерода по реакциям 3-4: ",
    f"RSH + NaOH = RSNa + H2O",
    f"COS + H2O → CO2 + H2S",
    f"H2S + 2NaOH → Na2S + 2H2O",
    f"CO2 + 2NaOH → Na2CO3 + H2O",
    f"Очищенное щелочью от меркаптановых соединений сырье проходит далее вверх через отстойную зону экстрактора {None}, снабженную металлическим каплеотбойником."
]

text = content1

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

#-----------------------------------------------------------------------------------------------------------------------
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

new_section.orientation = WD_ORIENT.PORTRAIT

if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

ch_7 = head_counter.increment()

heading = doc.add_heading(f'{ch_7:.0f} УСЛОВИЯ ПРОВЕДЕНИЯ ПРОЦЕССА «ДЕМЕРУС»', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

table7_1 = table_counter.increment()

text = [f'',
        f''
        ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df7_1 = read_excel_data('database.xlsx', '7.1')
df7_1 = df7_1.fillna('')

header_text_first = f'Таблица {table7_1} – Условия проведения процесса «ДЕМЕРУС»'
header_text_next = f'Продолжение таблицы {table7_1} – Условия проведения процесса «ДЕМЕРУС»'


rows_per_page_first = 16  # Количество строк для первой таблицы
rows_per_page_next = 18  # Количество строк для следующих таблиц

total_rows = len(df7_1)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df7_1, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df7_1, start_row, end_row)
    start_row = end_row

#-----------------------------------------------------------------------------------------------------------------------
new_section = doc.add_section(WD_SECTION.NEW_PAGE)
new_section.orientation = WD_ORIENT.PORTRAIT

ch_8 = head_counter.increment()

heading = doc.add_heading(f'{ch_8:.0f} НОРМЫ РАСХОДА ОСНОВНЫХ И ВСПОМОГАТЕЛЬНЫХ МАТЕРИАЛОВ', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = [f'',
        f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df8_1 = read_excel_data('database.xlsx', '8.1')
df8_1 = df8_1.fillna('')

table8_1 = table_counter.increment()
header_text = f'Таблица {table8_1} – Нормы расхода химреагентов и катализаторов при демеркаптанизации СУГ'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 100  # Количество строк для следующих таблиц

total_rows = len(df8_1)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text)
add_table(doc, df8_1, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text)
    add_table(doc, df8_1, start_row, end_row)
    start_row = end_row

text = [f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df8_2 = read_excel_data('database.xlsx', '8.2')
df8_2 = df8_2.fillna('')

table8_2 = table_counter.increment()
header_text = f'Таблица {table8_2} – Эксплуатационные расходы энергоресурсов при демеркаптанизации СУГ'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 100  # Количество строк для следующих таблиц

total_rows = len(df8_2)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text)
add_table(doc, df8_2, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text)
    add_table(doc, df8_2, start_row, end_row)
    start_row = end_row
#-----------------------------------------------------------------------------------------------------------------------
# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

ch_9 = head_counter.increment()

heading = doc.add_heading(f'{ch_9} МАТЕРИАЛЬНЫЙ И ТЕПЛОВОЙ БАЛАНС ТЕХНОЛОГИИ «ДЕМЕРУС»', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

ch_9_par_1 = par_counter.increment()

content1 = [
    "",
    "",
    f"{ch_9_par_1:.2} Исходные данные для расчета материального баланса",
    "",
    f"1. Число часов работы в году: {None}",
    f"2. Содержание сернистых соединений в СУГ до очистки:",
    f"- сероводород {None}, % масс.",
    f"- метилмеркаптан по сере {None}, % масс.",
    f"- этилмеркаптан по сере {None}, % масс.",
    f"- пропилмеркаптан по сере {None}, % масс.",
    f"- карбонилсульфид по сере {None}, % масс.",
    f"- сероуглерод {None}, % масс.",
    f"- общая меркаптановая сера {None}, % масс.",
    f"3. Содержание сернистых соединений в СУГ после очистки:",
    f"- сероводород {None}, % масс.",
    f"- метилмеркаптан {None}, % масс.",
    f"- этилмеркаптан {None}, % масс.",
    f"- пропилмеркаптан по сере {None}, % масс.",
    f"- карбонилсульфид по сере {None}, % масс.",
    f"- сероуглерод по сере {None}, % масс.",
    f"- общее содержание меркаптановой серы, не более {None}",
    f"- содержание диметилдисульфида по сере {None}",
    f"- содержание диэтилдисульфида по сере {None}",
    f"- содержание общей серы (не более) {None}",
    f"4. Диапазон устойчивой работы установки {None}%÷{None}%",
    f"Расчет материального баланса произведен из условий работы установки по сырью, равного {None} или {None} кг/ч."
]

ch_9_par_2 = par_counter.increment()

content2 = [
    "",
    f"{ch_9_par_2:.2} Расчетный материальный баланс технологии «ДЕМЕРУС»",
    "",
    f"Результаты расчета материального баланса технологии «ДЕМЕРУС» представлены в табл. {None}, схема материальных потоков блока очистки СУГ по технологии «ДЕМЕРУС» - на рис. {None}.",
    "",
    f"Расчет частоты замены 5% раствора щелочи в емкости {None}",
    "",
    f"Взаимодействие щелочи с углекислым газом протекает по следующей реакции:",
    f"NaOH + CO2 → NaHCO3;",
    f"Расход воздуха – {None} нм3/ч, содержание {None} – {None} % ({None} кг/ч).",
    f"Количество NaOH для полного превращения {None}:",
    f"{None} (кг/ч) ⸱ {None} (г/моль) / {None} (г/моль) = {None} (кг/ч),",
    f"где {None} – молярная масса {None}, {None} – молярная масса {None}.",
    f"Принимаем, что в реакцию вступит 4% NaOH из 5%:",
    f"{None} (кг/ч) ⸱ 5% / 4% = {None} (кг/ч)",
    f"Следовательно, необходимый расход NaOH – {None} кг/ч. Расход 5% раствора NaOH равен:",
    f"{None} / {None} = {None} кг/ч ⸱ {None} = {None} кг/год ",
    f"Объем емкости {None} – {None}, степень заполнения – {None}. Следовательно, емкость {None} вмещает:",
    f"{None} ⸱ {None} = {None} м3 ({None} кг раствора NaOH).",
    f"Частота замены раствора для обеспечения необходимого расхода NaOH:",
    f"{None} / {None} ⁓ {None} раз в год."
]

ch_9_par_3 = par_counter.increment()

content3 = [
    "",
    f"{ch_9_par_3:.2} Тепловой баланс стадии экстракции меркаптанов в {None}",
    "",
    f"Тепловой эффект взаимодействия этилмеркаптана со щелочью составляет {None} кДж/моль [14]. Для расчетов примем тепловой эффект взаимодействия метилмеркаптана со щелочью равным тепловому эффекту взаимодействия этилмеркаптана со щелочью.",
    f"Общее число молей меркаптанов, вступивших в реакцию со щелочью, составляет {None} моль/час:",
    f"{None} кг/ч ⸱ ({None}/{None}) (кг/моль) = {None} (моль/час)",
    f"1. Тепловой эффект реакции взаимодействия меркаптанов со щелочью равен:",
    f"Q1 = {None} ⸱ {None} = {None} кДж/час",
    f"Теплота реакции взаимодействия сероводорода со щелочью составляет – {None} кДж/моль. Общее число молей сероводорода, вступивших в реакцию со щелочью, составляет: {None} моль/час ({None} кг/ч⸱({None}/{None})(кг/моль) = {None}).",
    f"Q2 = {None} ⸱ {None} = {None} кДж/час",
    f"3. Количество теплоты, приходящее с СУГ:",
    f"Q3 = {None} ⸱ {None} = {None} кДж/час",
    f"где: {None} – уд. теплоемкость СУГ, кДж/кг °С;",
    f"4. Количество теплоты, приходящее со щелочным раствором:",
    f"Q4 = {None} ⸱ {None} = {None} кДж/час",
    f"где: {None} – уд. теплоемкость 10%-ного раствора щелочи, кДж/кг °С.",
    f"5. Теплопотери в колонне {None} Q5 принимаем по нормируемой плотности теплового потока qL в соответствии СП на тепловую изоляцию:",
    f"qL= {None} Вт/м;",
    f"Q5= {None} Вт/м ⸱ {None} м = {None} Вт = {None} кДж/ч",
    f"6. Количество теплоты, уносимое с очищенным СУГ:",
    f"Q6 = {None} ⸱ tвых ⸱ {None} = {None} ⸱ tвых",
    f"7. Количество теплоты, уносимое с насыщенным меркаптидами щелочным раствором:",
    f"Q7 = {None} ⸱ tвых ⸱ {None} = {None} ⸱ tвых",
    f"Решая уравнение Q1 + Q2 + Q3 + Q4 - Q5 = Q6 + Q7 относительно tвых, получим:",
    f"tвых= ({None} + {None} + {None} + {None} - {None}) / ({None}  + {None}) = {None}°С.",
    f"Следовательно, при смешении раствора щелочи с температурой {None}°С с СУГ, имеющим температуру {None}°С, температура СУГ в колонне {None} c учетом теплопотерь повысится на {None}℃ до {None} °С."
]


text = content1 + content2 + content3

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

#-----------------------------------------------------------------------------------------------------------------------
# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width


table9_1 = table_counter.increment()

def generate_flow_tables(df, flow_names, flow_nums):
    name_comp = df.iloc[:, 0]
    table_name_comp = ['Номер потока по схеме', 'Наименование потока', 'Состав'] + name_comp.tolist()

    flow_tables = []
    for i in range(len(flow_names)):
        flow_rate = df.iloc[:, 2 + i * 2].tolist()
        flow_frac = df.iloc[:, 3 + i * 2].tolist()

        table_flow_rate = [flow_nums[i], flow_names[i], 'кг/ч'] + [f'{item:.4f}' for item in flow_rate]
        table_flow_frac = ['', '', '% масс.'] + [f'{item:.4f}' for item in flow_frac]

        flow_tables.append((table_flow_rate, table_flow_frac))

    return table_name_comp, flow_tables


# Чтение данных из Excel-файла
df = read_excel_data('database.xlsx', 'Сводный мат баланс')
fn = read_excel_data('database.xlsx', 'Название потоков')

if df is not None and fn is not None:
    flow_names = fn.iloc[:, 1].tolist()
    flow_nums = fn.iloc[:, 0].tolist()

    table_name_comp, flow_tables = generate_flow_tables(df, flow_names, flow_nums)

    paragraph_after_break = doc.add_paragraph(
        f'Таблица {table9_1} – Расчетный материальный баланс процесса очистки СУГ')
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                         line_spacing=22, space_after=0, space_before=0)
    columns = [table_name_comp]
    for i in range(len(flow_tables)):
        columns.append(flow_tables[i][0])  # Номер потока
        columns.append(flow_tables[i][1])  # Доля потока

    len_status = all(len(col) == len(columns[0]) for col in columns)
    print(f"Все колонки имеют одинаковое количество элементов:", len_status)

    page_size = 6
    for start in range(1, len(columns), page_size):
        end = start + page_size
        current_columns = [table_name_comp] + columns[start:end]

        if len(current_columns) < page_size + 1:
            for _ in range(page_size + 1 - len(current_columns)):
                current_columns.append([''] * len(table_name_comp))

        table = doc.add_table(rows=len(table_name_comp), cols=len(current_columns))
        table.style = 'Table Grid'

        for row_idx in range(len(table_name_comp)):
            for col_idx, col_data in enumerate(current_columns):
                cell = table.cell(row_idx, col_idx)
                cell_paragraph = cell.paragraphs[0]
                cell_paragraph.text = str(col_data[row_idx])
                cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in cell_paragraph.runs:
                    set_font(run, 'Times New Roman', 12)
                set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                                     line_spacing=18, space_after=0, space_before=0)

        # Объединение ячеек
        for i in range(1, len(current_columns) - 1, 2):
            table.cell(0, i).merge(table.cell(0, i + 1))
            table.cell(1, i).merge(table.cell(1, i + 1))

        doc.add_page_break()

        paragraph_after_break = doc.add_paragraph(
            f'Продолжение таблицы {table9_1} – Расчетный материальный баланс процесса очистки СУГ')
        paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in paragraph_after_break.runs:
            set_font(run, 'Times New Roman', 14)
        set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                             line_spacing=22, space_after=0, space_before=0)


#-----------------------------------------------------------------------------------------------------------------------
table9_2 = table_counter.increment()

def generate_thermodynamics_tables(df, flow_names, flow_nums):
    thermodynamics_id = td.iloc[:, 0]
    table_thermodynamics_id = ['Номер потока по схеме', 'Наименование потока', 'Показатели'] + thermodynamics_id.tolist()

    thermodynamics_tables = []
    for i in range(len(flow_names)):
        liquid_phase = td.iloc[:, 2 + i * 2].tolist()
        gas_phase = td.iloc[:, 3 + i * 2].tolist()

        table_liquid_phase = [flow_nums[i], flow_names[i], 'Жидкая фаза'] + [f'{item:.4f}' for item in liquid_phase]
        table_gas_phase = ['', '', 'Газовая фаза'] + [f'{item:.4f}' for item in gas_phase]

        thermodynamics_tables.append((table_liquid_phase, table_gas_phase))

    return table_thermodynamics_id, thermodynamics_tables


# Чтение данных из Excel-файла
td = read_excel_data('database.xlsx', 'Термодинамика')
fn = read_excel_data('database.xlsx', 'Название потоков')

if td is not None and fn is not None:
    flow_names = fn.iloc[:, 1].tolist()
    flow_nums = fn.iloc[:, 0].tolist()

    table_thermodynamics_id, thermodynamics_tables = generate_thermodynamics_tables(td, flow_names, flow_nums)

    paragraph_after_break = doc.add_paragraph(
        f'Таблица {table9_2} – Термодинамические характеристики потоков процесса очистки СУГ')
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                         line_spacing=22, space_after=0, space_before=0)
    columns = [table_thermodynamics_id]
    for i in range(len(thermodynamics_tables)):
        columns.append(thermodynamics_tables[i][0])  # Номер потока
        columns.append(thermodynamics_tables[i][1])  # Доля потока

    len_status = all(len(col) == len(columns[0]) for col in columns)
    print(f"Все колонки имеют одинаковое количество элементов:", len_status)

    page_size = 6
    for start in range(1, len(columns), page_size):
        end = start + page_size
        current_columns = [table_thermodynamics_id] + columns[start:end]

        if len(current_columns) < page_size + 1:
            for _ in range(page_size + 1 - len(current_columns)):
                current_columns.append([''] * len(table_thermodynamics_id))

        table = doc.add_table(rows=len(table_thermodynamics_id), cols=len(current_columns))
        table.style = 'Table Grid'

        for row_idx in range(len(table_thermodynamics_id)):
            for col_idx, col_data in enumerate(current_columns):
                cell = table.cell(row_idx, col_idx)
                cell_paragraph = cell.paragraphs[0]
                cell_paragraph.text = str(col_data[row_idx])
                cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in cell_paragraph.runs:
                    set_font(run, 'Times New Roman', 12)
                set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                                     line_spacing=18, space_after=0, space_before=0)

        for i in range(1, len(current_columns) - 1, 2):
            table.cell(0, i).merge(table.cell(0, i + 1))
            table.cell(1, i).merge(table.cell(1, i + 1))

        doc.add_page_break()

        paragraph_after_break = doc.add_paragraph(
            f'Продолжение таблицы {table9_2} – Термодинамические характеристики потоков процесса очистки СУГ')
        paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in paragraph_after_break.runs:
            set_font(run, 'Times New Roman', 14)
        set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                             line_spacing=22, space_after=0, space_before=0)

#-----------------------------------------------------------------------------------------------------------------------
# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.LANDSCAPE

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width < new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width


fig9_1 = fig_counter.increment()

paragraph_after_break = doc.add_paragraph(
        f'Рисунок {fig9_1} – Материальные потоки блока очистки СУГ от меркаптанов')
paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in paragraph_after_break.runs:
    set_font(run, 'Times New Roman', 14)
set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                     line_spacing=22, space_after=0, space_before=0)

#-----------------------------------------------------------------------------------------------------------------------
# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.LANDSCAPE

# Убедимся, что размеры страницы корректны для альбомной ориентации
if new_section.page_width < new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

ch_10 = head_counter.increment()

heading = doc.add_heading(f'{ch_10} ФИЗИКО-ХИМИЧЕСКИЕ ОСНОВЫ ПРОЦЕССА', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

table10_1 = table_counter.increment()

text = [f'',
        f'',
        f'Таблица {table10_1:.1f} – Химизм процесса']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

#-----------------------------------------------------------------------------------------------------------------------
# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

ch_11 = head_counter.increment()

heading = doc.add_heading(f'{ch_11} СПЕЦИФИКАЦИЯ ОСНОВНОГО ТЕХНОЛОГИЧЕСКОГО ОБОРУДОВАНИЯ', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

ch_11_par_1 = par_counter.increment()

text = [f'',
        f'',
        f'{ch_11_par_1:.1f} Статическое оборудование']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

text = [f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df11_1 = read_excel_data('database.xlsx', '11.1')
df11_1 = df11_1.fillna('')

header_text_first = f'Таблица {ch_11_par_1:.1f} – Спецификация статического оборудования'
header_text_next = f'Продолжение таблицы {ch_11_par_1:.1f} – Спецификация статического оборудования'

rows_per_page_first = 9  # Количество строк для первой таблицы
rows_per_page_next = 12  # Количество строк для следующих таблиц

total_rows = len(df11_1)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df11_1, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df11_1, start_row, end_row)
    start_row = end_row

ch_11_par_2 = par_counter.increment()

text = [f'* давление верха/низа аппарата.',
        f'** уточняется на стадии детального проектирования',
        f'*** давление верха',
        f'',
        f'{ch_11_par_2} Теплообменное оборудование']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df11_2 = read_excel_data('database.xlsx', '11.2.1')
df11_2 = df11_2.fillna('')

header_text_first = f'Таблица {ch_11_par_2:.1f}.1 – Спецификация теплообменного оборудования по рабочей среде'
header_text_next = f'Продолжение таблицы {ch_11_par_2:.1f}.1 – Спецификация теплообменного оборудования по рабочей среде'

rows_per_page_first = 9  # Количество строк для первой таблицы
rows_per_page_next = 12  # Количество строк для следующих таблиц

total_rows = len(df11_2)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df11_2, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df11_2, start_row, end_row)
    start_row = end_row

text = [f'* уточняется на стадии детального проектирования.',
        f''
      ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

doc.add_page_break()

df11_3 = read_excel_data('database.xlsx', '11.2.2')
df11_3 = df11_3.fillna('')


header_text_first = f'Таблица {ch_11_par_2:.1f}.2 – Спецификация теплообменного оборудования по теплоносителю'
header_text_next = f'Продолжение таблицы {ch_11_par_2:.1f}.2 – Спецификация теплообменного оборудования по теплоносителю'

rows_per_page_first = 9  # Количество строк для первой таблицы
rows_per_page_next = 12  # Количество строк для следующих таблиц

total_rows = len(df11_3)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df11_3, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df11_3, start_row, end_row)
    start_row = end_row

ch_11_par_3 = par_counter.increment()

text = [f'* Уточняется при детальном проектировании и соответствует полной конденсации водяного пара.',
        f'',
        f'{ch_11_par_3:.1f} Динамическое оборудование'
      ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df11_3 = read_excel_data('database.xlsx', '11.3')
df11_3 = df11_3.fillna('')


header_text_first = f'Таблица {ch_11_par_3:.1f} – Спецификация динамического оборудования'
header_text_next = f'Продолжение таблицы {ch_11_par_3:.1f} – Спецификация динамического оборудования'

rows_per_page_first = 9  # Количество строк для первой таблицы
rows_per_page_next = 12  # Количество строк для следующих таблиц

total_rows = len(df11_3)
start_row = 0

# Первая таблица
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df11_3, start_row, end_row)
start_row = end_row

# Последующие таблицы
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df11_3, start_row, end_row)
    start_row = end_row

text = [f'* Уточняется на стадии детального проектирования.',
        f'** Предусмотреть резервную позицию на складе.'
        ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

doc.add_page_break()

# Сохраняем документ
doc.save('Мат баланс.docx')
